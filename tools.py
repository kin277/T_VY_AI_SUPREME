import os
import pandas as pd
import matplotlib.pyplot as plt
from github import Github
import zipfile
import secrets
import string
import smtplib
import requests
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import pytz
import pandas as pd
import requests
import qrcode
from pypdf import PdfReader
from deep_translator import GoogleTranslator
from PIL import Image
from docx import Document
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
from langchain_experimental.tools import PythonREPLTool
from langchain_core.tools import tool
from langchain_community.utilities.wolfram_alpha import WolframAlphaAPIWrapper
from langchain_community.tools import WolframAlphaQueryRun
from langchain_community.utilities import ArxivAPIWrapper
from langchain_community.tools import ArxivQueryRun
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_community.tools import WikipediaQueryRun
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_community.tools.tavily_search import TavilySearchResults

# 1. Python REPL
python_repl_tool = PythonREPLTool()

# 2. GitHub Search
@tool
def github_search_repos(query: str) -> str:
    """Tìm kiếm repository trên GitHub theo từ khóa query."""
    token = os.getenv("GITHUB_TOKEN", "")
    if not token:
        return "Chưa cấu hình GITHUB_TOKEN trong biến môi trường."
    try:
        g = Github(token)
        repos = g.search_repositories(query=query)
        results = [f"- {repo.full_name}: {repo.html_url}" for repo in repos[:5]]
        return "\n".join(results) if results else "Không tìm thấy repository nào."
    except Exception as e:
        return f"Lỗi GitHub API: {str(e)}"

# 3. Data Analysis
@tool
def analyze_data_file(file_path: str) -> str:
    """Đọc tệp CSV/Excel bằng pandas và trả về tóm tắt dữ liệu."""
    try:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        elif file_path.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file_path)
        else:
            return "Định dạng tệp không hỗ trợ (chỉ nhận .csv, .xls, .xlsx)."
        
        return (
            f"Tệp: {file_path}\n"
            f"Kích thước: {df.shape[0]} hàng, {df.shape[1]} cột\n"
            f"Danh sách cột: {list(df.columns)}\n\n"
            f"Mô tả thống kê:\n{df.describe().to_string()}"
        )
    except Exception as e:
        return f"Lỗi khi đọc tệp dữ liệu: {str(e)}"
        
# 1. Tải tệp từ URL bất kỳ về máy
@tool
def download_file_from_url(url: str, save_filename: str = "") -> str:
    """Tải tệp (ảnh, PDF, ZIP, MP3...) từ đường dẫn URL trên Internet về lưu tại máy."""
    try:
        response = requests.get(url, stream=True, timeout=15)
        response.raise_for_status()
        if not save_filename:
            save_filename = url.split("/")[-1].split("?")[0] or "downloaded_file"
        
        save_path = os.path.join(".", save_filename)
        with open(save_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        return f"Đã tải thành công tệp về địa chỉ: {save_path}"
    except Exception as e:
        return f"Lỗi khi tải tệp từ URL: {str(e)}"

# 2. Đọc và trích xuất nội dung văn bản từ file PDF
@tool
def read_pdf_file(file_path: str) -> str:
    """Đọc nội dung văn bản bên trong tệp PDF."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text[:4000] if text else "Tệp PDF trống hoặc là file ảnh quét (scanned)."
    except Exception as e:
        return f"Lỗi khi đọc tệp PDF: {str(e)}"

# 3. Tạo mã QR tự động
@tool
def generate_qr_code(data: str, filename: str = "qrcode.png") -> str:
    """Tạo ảnh mã QR từ văn bản hoặc đường link URL và lưu thành file PNG."""
    try:
        img = qrcode.make(data)
        img.save(filename)
        return f"Tạo mã QR thành công và đã lưu tại file: {filename}"
    except Exception as e:
        return f"Lỗi khi tạo mã QR: {str(e)}"

# 4. Tra cứu & Chuyển đổi tỷ giá tiền tệ
@tool
def convert_currency(from_curr: str, to_curr: str, amount: float) -> str:
    """Chuyển đổi tiền tệ giữa các quốc gia (Ví dụ: from_curr='USD', to_curr='VND', amount=100)."""
    try:
        url = f"https://open.er-api.com/v6/latest/{from_curr.upper()}"
        res = requests.get(url, timeout=5).json()
        if res.get("result") == "success":
            rates = res.get("rates", {})
            to_rate = rates.get(to_curr.upper())
            if to_rate:
                total = amount * to_rate
                return f"{amount:,.2f} {from_curr.upper()} = {total:,.2f} {to_curr.upper()}"
            return f"Không tìm thấy mã tiền tệ {to_curr}."
        return "Không thể kết nối đến máy chủ tỷ giá lúc này."
    except Exception as e:
        return f"Lỗi tra cứu tỷ giá: {str(e)}"

# 5. Dịch thuật văn bản đa ngôn ngữ
@tool
def translate_text(text: str, target_lang: str = "vi") -> str:
    """Dịch đoạn văn bản sang ngôn ngữ mong muốn (Mặc định 'vi' cho tiếng Việt, 'en' cho tiếng Anh)."""
    try:
        translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
        return translated
    except Exception as e:
        return f"Lỗi khi dịch thuật: {str(e)}"

# ==========================================
# 10 TIỆN ÍCH MỚI BỔ SUNG
# ==========================================

# 1. Tạo file Word (.docx)
@tool
def create_word_file(title: str, content: str, filename: str = "document.docx") -> str:
    """Tạo tệp Word (.docx) với tiêu đề và nội dung cho sẵn."""
    try:
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        if not filename.endswith('.docx'):
            filename += '.docx'
        doc.save(filename)
        return f"Đã tạo tệp Word thành công tại: {filename}"
    except Exception as e:
        return f"Lỗi khi tạo file Word: {str(e)}"

# 2. Tạo file Excel (.xlsx)
@tool
def create_excel_file(data_dict: dict, filename: str = "output.xlsx") -> str:
    """Tạo tệp Excel từ dữ liệu Dictionary (Ví dụ: {'Tên': ['An', 'Bình'], 'Tuổi': [20, 22]})."""
    try:
        df = pd.DataFrame(data_dict)
        if not filename.endswith('.xlsx'):
            filename += '.xlsx'
        df.to_excel(filename, index=False)
        return f"Đã xuất tệp Excel thành công tại: {filename}"
    except Exception as e:
        return f"Lỗi khi tạo file Excel: {str(e)}"

# 3. Gửi Email tự động qua SMTP
@tool
def send_email(to_email: str, subject: str, body: str) -> str:
    """Gửi email tự động. Yêu cầu biến môi trường SENDER_EMAIL và SENDER_PASSWORD."""
    sender_email = os.getenv("SENDER_EMAIL", "")
    sender_password = os.getenv("SENDER_PASSWORD", "")
    if not sender_email or not sender_password:
        return "Chưa cấu hình SENDER_EMAIL hoặc SENDER_PASSWORD trong biến môi trường."
    
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        return f"Đã gửi email thành công tới {to_email}"
    except Exception as e:
        return f"Lỗi khi gửi email: {str(e)}"

# 4. Rút gọn link URL
@tool
def shorten_url(long_url: str) -> str:
    """Rút gọn đường link URL dài bằng dịch vụ TinyURL."""
    try:
        api_url = f"http://tinyurl.com/api-create.php?url={long_url}"
        res = requests.get(api_url, timeout=5)
        return res.text if res.status_code == 200 else "Không thể rút gọn link."
    except Exception as e:
        return f"Lỗi khi rút gọn link: {str(e)}"

# 5. Tạo mật khẩu an toàn ngẫu nhiên
@tool
def generate_strong_password(length: int = 16) -> str:
    """Tạo mật khẩu ngẫu nhiên an toàn (bao gồm chữ hoa, chữ thường, số và ký tự đặc biệt)."""
    try:
        chars = string.ascii_letters + string.digits + "!@#$%^&*()"
        password = ''.join(secrets.choice(chars) for _ in range(max(length, 8)))
        return f"Mật khẩu ngẫu nhiên tạo ra: {password}"
    except Exception as e:
        return f"Lỗi khi tạo mật khẩu: {str(e)}"

# 6. Đọc nội dung văn bản từ trang web
@tool
def scrape_web_content(url: str) -> str:
    """Cào tiêu đề và các đoạn văn bản chính từ một địa chỉ trang web URL bất kỳ."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        res = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        for script in soup(["script", "style", "nav", "footer"]):
            script.extract()
            
        paragraphs = [p.get_text().strip() for p in soup.find_all('p') if p.get_text().strip()]
        content = "\n".join(paragraphs[:10])
        return f"Tiêu đề: {soup.title.string if soup.title else 'Không có'}\n\nNội dung:\n{content[:3000]}"
    except Exception as e:
        return f"Lỗi khi đọc trang web: {str(e)}"

# 7. Lấy lời thoại/phụ đề Video YouTube
@tool
def get_youtube_transcript(video_id_or_url: str) -> str:
    """Lấy nội dung phụ đề/lời thoại từ một video YouTube (nhập ID video hoặc link YouTube)."""
    try:
        video_id = video_id_or_url.split("v=")[-1].split("&")[0].split("/")[-1]
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['vi', 'en'])
        text = " ".join([t['text'] for t in transcript])
        return text[:3500]
    except Exception as e:
        return f"Không thể lấy phụ đề YouTube: {str(e)}"

# 8. Quy đổi múi giờ thế giới
@tool
def convert_timezone(time_str: str, from_tz: str, to_tz: str) -> str:
    """Chuyển đổi thời gian giữa các múi giờ (Ví dụ: time_str='2026-08-04 10:00', from_tz='Asia/Ho_Chi_Minh', to_tz='America/New_York')."""
    try:
        fmt = "%Y-%m-%d %H:%M"
        loc_from = pytz.timezone(from_tz)
        loc_to = pytz.timezone(to_tz)
        
        naive_dt = datetime.strptime(time_str, fmt)
        local_dt = loc_from.localize(naive_dt)
        target_dt = local_dt.astimezone(loc_to)
        
        return f"{time_str} ({from_tz}) = {target_dt.strftime(fmt)} ({to_tz})"
    except Exception as e:
        return f"Lỗi quy đổi múi giờ: {str(e)}"

# 9. Nén nhiều file thành 1 file ZIP
@tool
def zip_files(file_paths: str, zip_name: str = "archive.zip") -> str:
    """Nén các tệp lại thành file .zip (nhập danh sách file phân cách bởi dấu phẩy, ví dụ: 'a.txt, b.csv')."""
    try:
        files = [f.strip() for f in file_paths.split(",")]
        if not zip_name.endswith('.zip'):
            zip_name += '.zip'
            
        with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for f in files:
                if os.path.exists(f):
                    zipf.write(f, os.path.basename(f))
        return f"Đã nén thành công {len(files)} tệp vào {zip_name}"
    except Exception as e:
        return f"Lỗi khi nén file ZIP: {str(e)}"

# 10. Thay đổi kích thước ảnh (Resize Image)
@tool
def resize_image(image_path: str, width: int, height: int) -> str:
    """Thay đổi kích thước chiều rộng (width) và chiều cao (height) của một tệp ảnh."""
    try:
        with Image.open(image_path) as img:
            resized_img = img.resize((width, height))
            save_path = f"resized_{os.path.basename(image_path)}"
            resized_img.save(save_path)
        return f"Đã thay đổi kích thước ảnh và lưu tại: {save_path}"
    except Exception as e:
        return f"Lỗi khi xử lý ảnh: {str(e)}"
        
# 4. Wolfram Alpha
wolfram_tool = WolframAlphaQueryRun(api_wrapper=WolframAlphaAPIWrapper())

# 5. ArXiv
arxiv_tool = ArxivQueryRun(api_wrapper=ArxivAPIWrapper())

# 6. Wikipedia
wikipedia_tool = WikipediaQueryRun(api_wrapper=WikipediaAPIWrapper())

# 7. DuckDuckGo
duckduckgo_tool = DuckDuckGoSearchRun()

# 8. Tavily
tavily_tool = TavilySearchResults(max_results=3)

# DANH SÁCH TẤT CẢ TOOLS ĐỂ EXPORT
all_tools = [
    python_repl_tool,
    github_search_repos,
    analyze_data_file,
    wolfram_tool,
    arxiv_tool,
    wikipedia_tool,
    duckduckgo_tool,
    tavily_tool,
    download_file_from_url,
    read_pdf_file,
    generate_qr_code,
    convert_currency,
    translate_text,
    # 10 TIỆN ÍCH MỚI:
    create_word_file,
    create_excel_file,
    send_email,
    shorten_url,
    generate_strong_password,
    scrape_web_content,
    get_youtube_transcript,
    convert_timezone,
    zip_files,
    resize_image
]